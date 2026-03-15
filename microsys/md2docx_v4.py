import sys
import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_paragraph_style(p, is_rtl):
    """Unified setting of direction and alignment using both library and XML methods."""
    pPr = p._p.get_or_add_pPr()
    
    # 1. Set Direction (Bidi)
    for old_bidi in pPr.xpath('w:bidi'):
        pPr.remove(old_bidi)
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1' if is_rtl else '0')
    pPr.append(bidi)
    
    # 2. Set Alignment (Justification)
    for old_jc in pPr.xpath('w:jc'):
        pPr.remove(old_jc)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'right' if is_rtl else 'left')
    pPr.append(jc)

    # 3. Set Library property
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_rtl else WD_ALIGN_PARAGRAPH.LEFT

def set_run_font(run, font_name="Dubai", fallback="Noto Sans Arabic"):
    """Sets the primary and fallback fonts for a run."""
    run.font.name = font_name
    
    # Set Word-specific fonts via XML for better compatibility
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    
    # Second preference (fallback) often handled by system or Word themes, 
    # but we force the name explicitly in 'ascii' and 'hAnsi'.
    # If font_name is not found, Word often defaults to Noto if it's the next best.
    rPr.append(rFonts)

def set_run_rtl(run):
    """Sets run text direction to RTL."""
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    rPr.append(rtl)

def set_run_ltr(run):
    """Sets run text direction to LTR."""
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '0')
    rPr.append(rtl)

def set_table_rtl(table):
    """Sets table layout to RTL."""
    tblPr = table._tbl.tblPr
    bidiVisual = OxmlElement('w:bidiVisual')
    bidiVisual.set(qn('w:val'), '1')
    tblPr.append(bidiVisual)

def set_cell_background(cell, color_hex):
    """Sets the background color of a cell."""
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shd)

def set_cell_vertical_alignment(cell, align='center'):
    """Sets vertical alignment of cell content."""
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)

def set_cell_borders(cell, **kwargs):
    """Sets cell borders."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        if edge in kwargs:
            tag = OxmlElement(f'w:{edge}')
            for key, val in kwargs[edge].items():
                tag.set(qn(f'w:{key}'), str(val))
            tcBorders.append(tag)
    tcPr.append(tcBorders)

def set_keep_with_next(p):
    """Prevents paragraph from being separated from the next one on page breaks."""
    pPr = p._p.get_or_add_pPr()
    keepNext = OxmlElement('w:keepNext')
    keepNext.set(qn('w:val'), '1')
    pPr.append(keepNext)

def set_repeat_header(row):
    """Sets table row to repeat at the top of each page."""
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), '1')
    trPr.append(tblHeader)

def set_code_block_style(p):
    """Applies a code block style."""
    set_paragraph_style(p, is_rtl=False)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    for border_name in ['top', 'left', 'bottom', 'right']:
        edge = OxmlElement(f'w:{border_name}')
        edge.set(qn('w:val'), 'single')
        edge.set(qn('w:sz'), '4') 
        edge.set(qn('w:space'), '4')
        edge.set(qn('w:color'), 'D0D0D0')
        pbdr.append(edge)
    pPr.append(pbdr)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F4F4F4')
    pPr.append(shd)

def add_code_run(p, text, is_command=False):
    """Adds a run with monospace font for code."""
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    if is_command:
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 128)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Consolas')
    rFonts.set(qn('w:hAnsi'), 'Consolas')
    rPr.append(rFonts)
    set_run_ltr(run)

def parse_inline_styles(paragraph, text, is_rtl=True):
    """Parses bold text in markdown and adds to paragraph."""
    set_paragraph_style(paragraph, is_rtl)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            set_run_font(run)
            if is_rtl: set_run_rtl(run)
            else: set_run_ltr(run)
        else:
            run = paragraph.add_run(part)
            set_run_font(run)
            if is_rtl: set_run_rtl(run)
            else: set_run_ltr(run)

def detect_is_rtl(lines):
    """Detects if the document should be RTL."""
    arabic_pattern = re.compile(r'[\u0600-\u06FF]')
    for line in lines[:20]:
        if arabic_pattern.search(line):
            return True
    return False

def convert_md_to_docx(md_path, docx_path):
    doc = Document()
    try:
        with open(md_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
    except Exception as e:
        print(f"Error reading file: {e}")
        return

    is_bit_rtl = detect_is_rtl(lines)

    in_table = False
    table_data = []
    in_code_block = False

    def flush_table():
        nonlocal in_table, table_data
        if in_table and table_data:
            rows_count = len(table_data)
            cols_count = max(len(row) for row in table_data) if table_data else 0
            if rows_count > 0 and cols_count > 0:
                table = doc.add_table(rows=rows_count, cols=cols_count)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.RIGHT if is_bit_rtl else WD_TABLE_ALIGNMENT.LEFT
                if is_bit_rtl: set_table_rtl(table)
                set_repeat_header(table.rows[0])

                for i, row_data in enumerate(table_data):
                    table_cells = table.rows[i].cells
                    for j, cell_text in enumerate(row_data):
                        if j < len(table_cells):
                            cell = table_cells[j]
                            p = cell.paragraphs[0]
                            set_paragraph_style(p, is_bit_rtl)
                            set_cell_vertical_alignment(cell, 'center')
                            if i <= 2: set_keep_with_next(p)
                            
                            border_color = "DDDDDD"
                            border_style = {"sz": 4, "val": "single", "color": border_color}
                            cell_borders = {"top": border_style, "left": border_style, "bottom": border_style, "right": border_style}
                            
                            if i == 0: 
                                set_cell_background(cell, "E1F5FE")
                                cell_borders["bottom"] = {"sz": 12, "val": "single", "color": "B0BEC5"}
                            elif i % 2 == 0: 
                                set_cell_background(cell, "FAFAFA")
                            
                            set_cell_borders(cell, **cell_borders)
                            cleaned_text = cell_text.strip()
                            if cleaned_text.startswith('**') and cleaned_text.endswith('**'):
                                run = p.add_run(cleaned_text[2:-2]); run.bold = True
                            else:
                                run = p.add_run(cleaned_text)
                                if i == 0: run.bold = True
                            
                            set_run_font(run)
                            if is_bit_rtl: set_run_rtl(run)
                            else: set_run_ltr(run)
            table_data = []
        in_table = False

    for line in lines:
        raw_line = line.rstrip('\n')
        line = raw_line.strip()
        
        # Handle Images
        img_match = re.search(r'!\[(.*?)\]\((.*?)\)', line)
        if img_match:
            img_path = img_match.group(2)
            if os.path.exists(img_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, width=Inches(5))
                continue
            else:
                print(f"Warning: Image file not found at {img_path}")

        if line.startswith('```'):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            p = doc.add_paragraph(); set_code_block_style(p)
            is_cmd = line.startswith('$ ') or line.startswith('# ')
            add_code_run(p, raw_line, is_command=is_cmd)
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
            continue
            
        line = re.sub(r'<\|ref\|>.*?<\|/ref\|>', '', line)
        line = re.sub(r'<\|det\|>.*?<\|/det\|>', '', line)
        line = re.sub(r'--- Page \d+ ---', '', line)
        
        if line.startswith('|') and line.endswith('|'):
            in_table = True
            if '---' in line: continue
            table_data.append([cell.strip() for cell in line.strip('|').split('|')])
            continue
        elif in_table: flush_table()
        
        if not line: continue
        if line.startswith('<div') or line.startswith('</div') or line.startswith('<br'): continue
        if line == '<!-- pagebreak -->': doc.add_page_break(); continue
        if line == '---':
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_style(p, is_bit_rtl)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('________________________________________'); run.bold = True
            set_run_font(run)
            if is_bit_rtl: set_run_rtl(run)
            else: set_run_ltr(run)
            p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(12)
            continue
            
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            title_text = line.lstrip('#').strip().replace('**', '')
            p = doc.add_heading('', level=min(level, 9))
            set_paragraph_style(p, is_bit_rtl)
            run = p.add_run(title_text)
            set_run_font(run)
            if is_bit_rtl: set_run_rtl(run)
            else: set_run_ltr(run)
            continue
            
        if line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            parse_inline_styles(p, line[2:].strip(), is_rtl=is_bit_rtl)
            continue
            
        p = doc.add_paragraph()
        parse_inline_styles(p, line, is_rtl=is_bit_rtl)
        
    flush_table()
    try:
        doc.save(docx_path)
        print(f"Successfully saved to {docx_path}")
    except Exception as e:
        print(f"Error saving to {docx_path}: {e}")

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python md2docx_v4.py [input] [output]"); sys.exit(1)
    convert_md_to_docx(sys.argv[1], sys.argv[2])
